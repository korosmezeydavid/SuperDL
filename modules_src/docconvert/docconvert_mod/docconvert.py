"""Dokumentum-konverter – azonnali konverzió szöveg-, könyv- és KÉP-formátumok
közt, a lehető legjobb minőségű úton, motoros routinggal:

  • KÉP (TIF/PNG/JPG…) → OCR (AI-vízió vagy Tesseract) → szöveg → bármely kimenet
  • MOBI / AZW3 / PDF-kimenet, illetve MOBI / DOC bemenet → Calibre (ebook-convert)
    vagy LibreOffice, ha telepítve van
  • RTF / ODT / Markdown / FB2 → Pandoc (igény szerint letöltve, ~/.superdl/bin)
  • a többi (TXT, DOCX, EPUB, PDF-bemenet, HTML) → a beágyazott libekkel

KÓDOLÁS:  UTF-8, UTF-8 BOM-mal, UTF-16, Windows-1250, ISO-8859-2, Windows-1252
"""

import html as _html
import os
import subprocess
import tempfile
from pathlib import Path

# MODUL: a megosztott segédek a Core futtatókörnyezetéből (a superdl csomagból),
# a converter-specifikus kód viszont a modulban van
from superdl import booktext
from superdl import extratools
from superdl import ocr

_NOWIN = 0x08000000 if os.name == "nt" else 0

# (megjelenített név, kiterjesztés, kell-e külső eszköz)
OUT_FORMATS = [
    ("Szöveg (TXT)", "txt", None),
    ("Weblap (HTML)", "html", None),
    ("Word (DOCX)", "docx", None),
    ("E-könyv (EPUB)", "epub", None),
    ("Rich Text (RTF)", "rtf", "pandoc"),
    ("OpenDocument (ODT)", "odt", "pandoc"),
    ("Markdown (MD)", "md", "pandoc"),
    ("FictionBook (FB2)", "fb2", "pandoc"),
    ("Kindle (MOBI)", "mobi", "calibre"),
    ("Kindle (AZW3)", "azw3", "calibre"),
    ("PDF", "pdf", "calibre/office"),
]
ENCODINGS = [("Automatikus felismerés", "auto"),
             ("UTF-8", "utf-8"),
             ("UTF-8 BOM-mal", "utf-8-sig"),
             ("UTF-16", "utf-16"),
             ("Közép-európai (Windows-1250)", "cp1250"),
             ("Közép-európai (ISO-8859-2 / Latin-2)", "iso-8859-2"),
             ("Magyar DOS (CP852 / Latin-2)", "cp852"),
             ("Magyar CWI-2 (régi DOS)", "cwi2"),
             ("Régi DOS (CP437)", "cp437"),
             ("Nyugat-európai (Windows-1252)", "cp1252")]

# az AUTOMATIKUS felismerés egybájtos jelöltjei (a legvalószínűbbtől); a helyeset
# NEM az „elsőként hibátlanul dekódol" választja (az iso-8859-2/cp852/cp437 MINDEN
# bájtot elfogad, így kacatot is), hanem a `_decode_score` szerinti LEGJOBB.
_TRY_ENC = ["cp1250", "cp852", "iso-8859-2", "cp437", "cp1252"]

# a helyes magyar dekódolás erős jelei
_HU_CHARS = "áéíóöőúüűÁÉÍÓÖŐÚÜŰ"
# a magyar szövegben megszokott, nem-ASCII, de RENDBEN lévő írásjelek
_OK_EXTRA = set("„”“‚’‘–—…«»•·§°€£")

# CWI-2 (régi magyar DOS-kódlap, ≈CP3845): a CP437-re épül; a magyar kisbetűk
# (á é í ó ö ú ü) és az É Ö Ü nagybetűk MÁR a CP437 helyükön vannak, csak a
# ő ű Ő Ű és néhány nagybetű (Á Í Ó Ú) kerül más pozícióra. Ezért a CP437-alapot
# vesszük, és CSAK ezt a 8 magyar pozíciót írjuk felül → teljes magyar lefedettség.
_CWI2_OVERRIDES = {0x8D: "Í", 0x8F: "Á", 0x93: "ő", 0x95: "Ó",
                   0x96: "ű", 0x97: "Ú", 0x98: "Ű", 0xA7: "Ő"}
_CWI2_TABLE = list(bytes(range(256)).decode("cp437"))
for _b, _ch in _CWI2_OVERRIDES.items():
    _CWI2_TABLE[_b] = _ch
_CWI2_TABLE = "".join(_CWI2_TABLE)


def decode_cwi2(data: bytes) -> str:
    """A régi magyar CWI-2 DOS-kódlapú bájtok Unicode-ra fejtése."""
    return "".join(_CWI2_TABLE[b] for b in data)

# bemeneti kiterjesztések (a fájlválasztóhoz)
IN_EXTS = (".txt", ".docx", ".epub", ".pdf", ".html", ".htm", ".rtf", ".odt",
           ".md", ".markdown", ".fb2", ".doc", ".mobi", ".azw3"
           ) + ocr.IMAGE_EXTS

# ---- formátum-halmazok a routinghoz ----
PHASE1_WRITE = {"txt", "html", "docx", "epub"}
PANDOC_WRITE = {"rtf", "odt", "md", "markdown", "fb2"}
PANDOC_READ = {".rtf", ".odt", ".md", ".markdown", ".fb2", ".org",
               ".tex", ".latex", ".docx", ".epub", ".html", ".htm", ".txt"}
CALIBRE_OUT = {"mobi", "azw3", "pdf"}
CALIBRE_IN = {".mobi", ".azw3", ".azw", ".doc", ".lit", ".pdb"}


def _run(cmd, timeout=600):
    try:
        r = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                           text=True, encoding="utf-8", errors="replace",
                           creationflags=_NOWIN, timeout=timeout)
        return r.returncode, (r.stdout or "") + (r.stderr or "")
    except (OSError, subprocess.SubprocessError) as e:
        return 1, str(e)


# ---- bemeneti olvasók (phase-1, szöveg-kinyerés) ----------------------

def _decode_score(text: str) -> int:
    """A dekódolt szöveg „hihetőségének" pontszáma. A magyar ékezetes betűk
    ERŐSEN növelik, a vezérlő-/dobozrajz-/pótló-karakterek CSÖKKENTIK – így a
    régi DOS-kódlapok (cp852/cp437) helyesen felismerhetők a mindent elfogadó
    iso-8859-2 helyett."""
    good = bad = 0
    for ch in text:
        if ch in "\r\n\t":
            continue
        o = ord(ch)
        if ch in _HU_CHARS:             # magyar ékezetes – a jó kódlap erős jele
            good += 3
        elif o < 0x80:                  # ASCII
            if o < 32 or o == 0x7f:     # vezérlőkarakter
                bad += 5
            else:                       # normál ASCII betű/szám/írásjel
                good += 1
        elif ch in _OK_EXTRA:           # megszokott európai írásjel
            good += 1
        else:                           # bármi más nem-ASCII (nem magyar): gyanús
            bad += 2                    # (Š, Ł, ŕ, ˘, dobozrajz, pótló-karakter…)
    return good - bad


def _auto_decode(data: bytes) -> str:
    """A bájtokat a legvalószínűbb kódlappal fejti meg: előbb szigorú UTF-8,
    aztán a `_TRY_ENC` egybájtos jelöltek közül a LEGJOBB pontszámú."""
    for enc in ("utf-8-sig", "utf-8"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            pass
    best_text, best_score = None, None
    candidates = [(enc, None) for enc in _TRY_ENC] + [("cwi2", decode_cwi2)]
    for enc, fn in candidates:
        if fn is not None:
            t = fn(data)
        else:
            try:
                t = data.decode(enc)
            except UnicodeDecodeError:
                t = data.decode(enc, errors="replace")
        sc = _decode_score(t)
        if best_score is None or sc > best_score:
            best_text, best_score = t, sc
    return best_text if best_text is not None else data.decode("utf-8",
                                                               errors="replace")


def _read_txt(path: Path, in_encoding):
    data = path.read_bytes()
    if in_encoding and in_encoding != "auto":
        text = decode_cwi2(data) if in_encoding == "cwi2" \
            else data.decode(in_encoding, errors="replace")
    else:
        text = _auto_decode(data)
    return booktext.Book(title=path.stem, sections=[booktext._clean(text)])


def _read_html(path: Path):
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(path.read_bytes(), "html.parser")
    title = (soup.title.string.strip()
             if soup.title and soup.title.string else path.stem)
    for tag in soup(["script", "style"]):
        tag.decompose()
    return booktext.Book(title=title,
                         sections=[booktext._clean(soup.get_text("\n"))])


def read_document(path, in_encoding=None) -> booktext.Book:
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".txt":
        return _read_txt(p, in_encoding)
    if ext in (".html", ".htm"):
        return _read_html(p)
    if ext in (".docx", ".epub", ".pdf"):
        return booktext.extract(str(p))
    raise ValueError(f"Nem támogatott bemeneti formátum: {ext}")


# ---- phase-1 írók -----------------------------------------------------

def _lines(book):
    for sec in book.sections:
        for line in sec.split("\n"):
            line = line.strip()
            if line:
                yield line
        yield ""


def _write_txt(book, dst, out_encoding):
    dst.write_bytes(book.text.encode(out_encoding or "utf-8", errors="replace"))


def _write_html(book, dst):
    body = [f"<h1>{_html.escape(book.title)}</h1>"]
    for line in _lines(book):
        body.append(f"<p>{_html.escape(line)}</p>" if line else "")
    doc = ("<!DOCTYPE html>\n<html lang=\"hu\">\n<head>\n"
           "<meta charset=\"utf-8\">\n"
           f"<title>{_html.escape(book.title)}</title>\n</head>\n<body>\n"
           + "\n".join(body) + "\n</body>\n</html>\n")
    dst.write_text(doc, encoding="utf-8")


def _write_docx(book, dst):
    import docx
    d = docx.Document()
    d.add_heading(book.title, level=1)
    for i, sec in enumerate(book.sections):
        if i:
            d.add_paragraph("")
        for line in sec.split("\n"):
            if line.strip():
                d.add_paragraph(line.strip())
    d.save(str(dst))


def _write_epub(book, dst):
    from ebooklib import epub
    b = epub.EpubBook()
    b.set_title(book.title)
    b.set_language("hu")
    chapters = []
    for i, sec in enumerate(book.sections):
        paras = "".join(f"<p>{_html.escape(ln.strip())}</p>"
                        for ln in sec.split("\n") if ln.strip())
        c = epub.EpubHtml(title=f"{i + 1}. rész",
                          file_name=f"chap_{i + 1}.xhtml", lang="hu")
        c.content = f"<h2>{_html.escape(book.title)}</h2>{paras}"
        b.add_item(c)
        chapters.append(c)
    if not chapters:
        c = epub.EpubHtml(title="1. rész", file_name="chap_1.xhtml", lang="hu")
        c.content = f"<h2>{_html.escape(book.title)}</h2>"
        b.add_item(c)
        chapters = [c]
    b.toc = chapters
    b.add_item(epub.EpubNcx())
    b.add_item(epub.EpubNav())
    b.spine = ["nav"] + chapters
    epub.write_epub(str(dst), b)


def _unicode_font() -> str | None:
    """Egy Unicode TrueType betű a magyar ékezetekhez (a beépített PDF-hez). A
    Windowson mindig elérhető Arialt részesítjük előnyben, tartalékként a
    csomagolt DejaVuSans (ha mellécsomagoltuk)."""
    import sys as _sys
    win = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
    for name in ("arial.ttf", "segoeui.ttf", "tahoma.ttf", "verdana.ttf",
                 "calibri.ttf"):
        c = win / name
        if c.is_file():
            return str(c)
    mei = getattr(_sys, "_MEIPASS", None)
    if mei:
        d = Path(mei) / "DejaVuSans.ttf"
        if d.is_file():
            return str(d)
    return None


def _write_pdf(book, dst):
    """Tiszta-Python PDF a kinyert szövegből (fpdf2) – KÜLSŐ PROGRAM NÉLKÜL.
    A magyar ékezetekhez beágyazott Unicode TrueType betűt használ. Így a
    szöveg/DOCX/EPUB/HTML → PDF konverzióhoz NEM kell Calibre/LibreOffice."""
    from fpdf import FPDF
    font = _unicode_font()
    pdf = FPDF(format="A4")
    pdf.set_auto_page_break(True, margin=15)
    pdf.add_page()
    fam = "Helvetica"
    if font:
        pdf.add_font("doc", "", font)
        fam = "doc"
    pdf.set_font(fam, size=16)
    # új sor a bal margóra (különben a kurzor a jobb szélen marad → „nincs hely")
    pdf.multi_cell(0, 9, book.title or "", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)
    pdf.set_font(fam, size=12)
    for line in _lines(book):
        pdf.multi_cell(0, 7, line if line else " ",
                       new_x="LMARGIN", new_y="NEXT")
    pdf.output(str(dst))


def _write_simple(book, dst: Path, out_format, out_encoding):
    if out_format == "txt":
        _write_txt(book, dst, out_encoding)
    elif out_format == "html":
        _write_html(book, dst)
    elif out_format == "docx":
        _write_docx(book, dst)
    elif out_format == "epub":
        _write_epub(book, dst)
    else:
        raise ValueError(f"Belső hiba: ismeretlen egyszerű formátum {out_format}")


# ---- külső eszköz-hívók ----------------------------------------------

def _pandoc(pandoc, src, dst):
    rc, out = _run([pandoc, str(src), "-o", str(dst)])
    if rc != 0 or not Path(dst).exists():
        raise RuntimeError("Pandoc hiba: "
                           + (out.strip().splitlines()[-1] if out.strip()
                              else "ismeretlen"))


def _calibre(cal, src, dst):
    rc, out = _run([cal, str(src), str(dst)], timeout=900)
    if rc != 0 or not Path(dst).exists():
        raise RuntimeError("Calibre hiba: "
                           + (out.strip().splitlines()[-1] if out.strip()
                              else "ismeretlen"))


def _office(office, src, dst, out_format):
    d = Path(dst)
    rc, out = _run([office, "--headless", "--convert-to", out_format,
                    "--outdir", str(d.parent), str(src)], timeout=900)
    produced = d.parent / (Path(src).stem + "." + out_format)
    if rc != 0:
        raise RuntimeError("LibreOffice hiba: "
                           + (out.strip().splitlines()[-1] if out.strip()
                              else "ismeretlen"))
    if produced != d and produced.exists():
        produced.replace(d)
    if not d.exists():
        raise RuntimeError("A LibreOffice nem készítette el a fájlt.")


def _need_msg(out_format, in_ext) -> str:
    if out_format in ("mobi", "azw3"):
        return ("A MOBI/AZW3 készítéséhez a Calibre szükséges (ingyenes). "
                "Telepítsd a calibre-ebook.com oldalról, és a SuperDL "
                "automatikusan felismeri.")
    if out_format == "pdf" or in_ext == ".doc":
        return ("Ehhez a Calibre VAGY a LibreOffice szükséges (mindkettő "
                "ingyenes). Telepítsd egyiket, és a SuperDL felismeri.")
    return "Ehhez a konverzióhoz egy külső eszköz (Calibre/LibreOffice) kell."


# ---- könyv → bármely kimenet -----------------------------------------

def _write_any(book, dst, out_format, out_encoding, progress=None):
    """Egy kinyert Book kiírása a kívánt formátumba – ha kell, Pandoc/Calibre
    közbeiktatásával (a könyvet egy köztes HTML/EPUB-on át)."""
    d = Path(dst)
    if out_format in PHASE1_WRITE:
        _write_simple(book, d, out_format, out_encoding)
    elif out_format in PANDOC_WRITE:
        pandoc = extratools.ensure_pandoc(progress)
        if not pandoc:
            raise RuntimeError("Ehhez a kimenethez Pandoc kell (az első "
                               "használathoz internet szükséges).")
        tmp = Path(tempfile.gettempdir()) / (d.stem + "_tmp.html")
        _write_html(book, tmp)
        try:
            _pandoc(pandoc, tmp, d)
        finally:
            tmp.unlink(missing_ok=True)
    elif out_format in CALIBRE_OUT:
        tmp = Path(tempfile.gettempdir()) / (d.stem + "_tmp.epub")
        _write_epub(book, tmp)
        try:
            cal = extratools.find_calibre()
            office = extratools.find_libreoffice()
            if cal:
                _calibre(cal, tmp, d)
            elif office and out_format == "pdf":
                _office(office, tmp, d, "pdf")
            else:
                raise RuntimeError(_need_msg(out_format, ""))
        finally:
            tmp.unlink(missing_ok=True)
    else:
        raise ValueError(f"Nem támogatott kimeneti formátum: {out_format}")


# ---- a fő belépő ------------------------------------------------------

def convert(src, dst, out_format, in_encoding=None, out_encoding="utf-8",
            ocr_engine="ai", progress=None) -> str:
    """Konvertálás a legjobb elérhető úton. Visszaad: felolvasható eredmény."""
    in_ext = Path(src).suffix.lower()

    # 1) KÉP → OCR → szöveg → kimenet
    if in_ext in ocr.IMAGE_EXTS:
        text = ocr.ocr(src, ocr_engine)
        if not text.strip():
            raise RuntimeError("Az OCR nem talált szöveget a képen.")
        book = booktext.Book(title=Path(src).stem,
                             sections=[booktext._clean(text)])
        _write_any(book, dst, out_format, out_encoding, progress)
        eng = ocr.ENGINES.get(ocr_engine, ocr_engine)
        return (f"OCR kész ({eng}) → {Path(dst).name}. "
                f"{len(text)} karakter felismerve.")

    # 2) MOBI/AZW3/PDF-kimenet vagy MOBI/DOC-bemenet → Calibre/LibreOffice (direkt)
    if out_format in CALIBRE_OUT or in_ext in CALIBRE_IN:
        cal = extratools.find_calibre()
        if cal:
            _calibre(cal, src, dst)
            return f"Kész (Calibre) → {Path(dst).name}."
        office = extratools.find_libreoffice()
        if office and (out_format == "pdf" or in_ext == ".doc"):
            _office(office, src, dst, out_format)
            return f"Kész (LibreOffice) → {Path(dst).name}."
        # PDF KÜLSŐ PROGRAM NÉLKÜL is: a kinyert szövegből beépített (fpdf2)
        # PDF-et készítünk → a szöveg/DOCX/EPUB/HTML/PDF→PDF sosem zsákutca
        # (egy felhasználó jelezte, hogy nincs Calibre/LibreOffice-a, mégis
        # PDF-et szeretne)
        if out_format == "pdf" and in_ext not in CALIBRE_IN:
            try:
                book = read_document(src, in_encoding)
            except ValueError:
                book = None
            if book is not None:
                _write_pdf(book, dst)
                return (f"Kész (beépített PDF) → {Path(dst).name}. "
                        f"{book.chars} karakter – külső program nem kellett.")
        raise RuntimeError(_need_msg(out_format, in_ext))

    # 3) RTF/ODT/MD/FB2 érintett → Pandoc DIREKT (legjobb formázás-megőrzés)
    pandoc_relevant = (out_format in PANDOC_WRITE
                       or in_ext in {".rtf", ".odt", ".md", ".markdown",
                                     ".fb2", ".org", ".tex", ".latex"})
    if pandoc_relevant and in_ext in PANDOC_READ:
        pandoc = extratools.ensure_pandoc(progress)
        if not pandoc:
            raise RuntimeError("A Pandoc nem érhető el, és nem sikerült "
                               "letölteni (az első használathoz internet kell).")
        _pandoc(pandoc, src, dst)
        return f"Kész (Pandoc) → {Path(dst).name}."

    # 4) phase-1: kinyerés a beágyazott libekkel → kiírás (ha kell, Pandoc/Calibre)
    book = read_document(src, in_encoding)
    _write_any(book, dst, out_format, out_encoding, progress)
    enc = f", {out_encoding} kódolással" if out_format == "txt" else ""
    return (f"Kész: „{book.title}” → {Path(dst).name} ({out_format.upper()}"
            f"{enc}). {book.chars} karakter, {len(book.sections)} szakasz.")
